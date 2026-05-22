"""
File parser — extracts plain text from uploaded documents.

Supports PDF (.pdf), Word (.docx), Markdown (.md), and plain text
(.txt, .text) formats.  Old binary Word files (.doc) are explicitly
rejected with a helpful error message.

This module is intentionally synchronous so callers can run it in a
thread via ``asyncio.to_thread`` when needed.
"""

from __future__ import annotations

import io
import logging
from dataclasses import dataclass, field
from pathlib import PurePosixPath

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Public constants
# ---------------------------------------------------------------------------

SUPPORTED_EXTENSIONS: set[str] = {".pdf", ".docx", ".doc", ".md", ".txt", ".text"}

MAX_FILE_SIZE_BYTES: int = 20 * 1024 * 1024  # 20 MB


# ---------------------------------------------------------------------------
# Data container
# ---------------------------------------------------------------------------

@dataclass
class ParsedDocument:
    """Result of parsing a file into plain text.

    Attributes
    ----------
    text:
        Extracted plain-text content.
    page_count:
        Number of pages for paginated formats (PDF); ``1`` for everything else.
    metadata:
        Optional metadata extracted from the document (e.g. title, author).
    """

    text: str
    page_count: int = 1
    metadata: dict[str, str] = field(default_factory=dict)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def parse_file(file_bytes: bytes, filename: str) -> ParsedDocument:
    """Parse *file_bytes* based on *filename*'s extension and return a
    :class:`ParsedDocument`.

    Parameters
    ----------
    file_bytes:
        Raw bytes of the uploaded file.
    filename:
        Original filename (used to determine format by extension).

    Raises
    ------
    ValueError
        If the file exceeds :data:`MAX_FILE_SIZE_BYTES`, has an
        unsupported extension, or cannot be parsed.
    """
    # --- Size guard ---
    if len(file_bytes) > MAX_FILE_SIZE_BYTES:
        raise ValueError(
            f"File is too large ({len(file_bytes) / (1024 * 1024):.1f} MB). "
            f"Maximum allowed size is {MAX_FILE_SIZE_BYTES / (1024 * 1024):.0f} MB."
        )

    ext = PurePosixPath(filename).suffix.lower()

    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file format '{ext}'. "
            f"Supported formats: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    if ext == ".pdf":
        result = _parse_pdf(file_bytes, filename)
    elif ext == ".docx":
        result = _parse_docx(file_bytes, filename)
    elif ext == ".doc":
        raise ValueError(
            "The legacy .doc format (binary Word) is not supported. "
            "Please convert the file to .docx and re-upload."
        )
    else:
        # .md, .txt, .text
        result = _parse_text(file_bytes, filename)

    # Sanitize text — remove null bytes and control chars that break PostgreSQL
    result.text = _sanitize_text(result.text)
    return result


# ---------------------------------------------------------------------------
# Text sanitisation
# ---------------------------------------------------------------------------

def _sanitize_text(text: str) -> str:
    """Remove null bytes and non-printable control characters.

    PostgreSQL rejects strings containing ``\\x00``.  PDF extractors
    often emit null bytes, form-feed (``\\x0c``), and other C0/C1
    control codes when processing scanned or math-heavy documents.

    We keep common whitespace (tab ``\\t``, newline ``\\n``,
    carriage-return ``\\r``) and strip everything else below U+0020.
    """
    # Fast path
    if "\x00" not in text and all(ch >= " " or ch in "\t\n\r" for ch in text[:200]):
        return text

    cleaned = []
    for ch in text:
        if ch == "\x00":
            continue  # always strip null
        cp = ord(ch)
        # Keep tab, newline, carriage-return; drop other C0 controls (0x01-0x08, 0x0B-0x0C, 0x0E-0x1F)
        if cp < 0x20 and ch not in ("\t", "\n", "\r"):
            cleaned.append(" ")  # replace with space to preserve word boundaries
            continue
        cleaned.append(ch)

    result = "".join(cleaned)
    if len(result) < len(text):
        logger.info(
            "Sanitized text: removed %d problematic characters",
            len(text) - len(result),
        )
    return result


# ---------------------------------------------------------------------------
# Private parsers
# ---------------------------------------------------------------------------

def _parse_pdf(file_bytes: bytes, filename: str) -> ParsedDocument:
    """Extract text and metadata from a PDF file."""
    try:
        from pypdf import PdfReader  # type: ignore[import-untyped]
    except ImportError as exc:
        raise ImportError(
            "The 'pypdf' package is required to parse PDF files. "
            "Install it with: pip install pypdf"
        ) from exc

    reader = PdfReader(io.BytesIO(file_bytes))

    pages: list[str] = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)

    full_text = "\n\n".join(pages)
    page_count = len(reader.pages)

    # Extract available metadata
    metadata: dict[str, str] = {}
    info = reader.metadata
    if info:
        if info.title:
            metadata["title"] = str(info.title)
        if info.author:
            metadata["author"] = str(info.author)
        if info.subject:
            metadata["subject"] = str(info.subject)

    if not full_text.strip():
        logger.warning("PDF '%s' produced no extractable text (scanned image?)", filename)

    logger.info(
        "Parsed PDF '%s': %d pages, %d chars extracted",
        filename,
        page_count,
        len(full_text),
    )
    return ParsedDocument(text=full_text, page_count=page_count, metadata=metadata)


def _parse_docx(file_bytes: bytes, filename: str) -> ParsedDocument:
    """Extract text from a DOCX file (paragraphs and table cells)."""
    try:
        from docx import Document  # type: ignore[import-untyped]
    except ImportError as exc:
        raise ImportError(
            "The 'python-docx' package is required to parse DOCX files. "
            "Install it with: pip install python-docx"
        ) from exc

    doc = Document(io.BytesIO(file_bytes))

    parts: list[str] = []

    # Paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)

    full_text = "\n".join(parts)

    # Extract core properties as metadata
    metadata: dict[str, str] = {}
    core = doc.core_properties
    if core.title:
        metadata["title"] = str(core.title)
    if core.author:
        metadata["author"] = str(core.author)
    if core.subject:
        metadata["subject"] = str(core.subject)

    if not full_text.strip():
        logger.warning("DOCX '%s' produced no extractable text", filename)

    logger.info(
        "Parsed DOCX '%s': %d chars extracted",
        filename,
        len(full_text),
    )
    return ParsedDocument(text=full_text, page_count=1, metadata=metadata)


def _parse_text(file_bytes: bytes, filename: str) -> ParsedDocument:
    """Decode raw text files (UTF-8, with fallback for encoding errors)."""
    text = file_bytes.decode("utf-8", errors="replace")

    logger.info(
        "Parsed text file '%s': %d chars",
        filename,
        len(text),
    )
    return ParsedDocument(text=text, page_count=1, metadata={})
